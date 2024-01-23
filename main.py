import docx
from docx.shared import Inches
import time
import OCR
from selenium import webdriver
from selenium.webdriver.common import alert
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions
from selenium.webdriver.support.wait import WebDriverWait

driver= webdriver.Chrome(executable_path="C:/Automation Testing/drivers/chromedriver.exe")
driver.get("https://tempmailo.com/")
value=driver.find_element_by_id("i-email").get_attribute('value')
with open('C:/Users/KumarGaurav/Desktop/AccountforTest.txt','a') as writer:
    if(value):
        writer.write(" Email id - "+value)
        writer.write("      ")
        writer.write("Password -  Password123@")
        writer.write('\n')

print(value)
driver.implicitly_wait(20)
driver.execute_script("window.open('');")
driver.switch_to.window(driver.window_handles[1])
driver.get("Enter website link to automate")


driver.find_element_by_xpath("//a[contains(text(),'MY ACCOUNT')]").click()
driver.find_element_by_link_text("Don't have an ID?").click()
driver.find_element_by_id("text1").send_keys(value)
driver.find_element_by_id("btnContinue").click()
time.sleep(2)
driver.save_screenshot('C:/Users/KumarGaurav/Desktop/screenTest/screen1.png')
driver.close()
driver.switch_to.window(driver.window_handles[0])

wait = WebDriverWait(driver,10)
driver.find_element_by_xpath("//span[text()='Refresh']").click()
driver.find_element_by_xpath("//span[text()='Refresh']").click()
driver.find_element_by_xpath("//span[text()='Refresh']").click()
driver.find_element_by_xpath("//span[text()='Refresh']").click()
driver.find_element_by_xpath('''//div[text()='"Link" <myaccount@email.com>']''').click()
driver.save_screenshot('C:/Users/KumarGaurav/Desktop/screenTest/screen2.png')
driver.switch_to.frame('fullmessage')

wait.until(expected_conditions.presence_of_element_located((By.XPATH, "//a[@target='_blank']")))
driver.find_element_by_xpath("//a[@target='_blank']").click()
driver.switch_to.window(driver.window_handles[1])
driver.find_element_by_id("text-firstName").send_keys("Tamesh")
driver.find_element_by_id("text-lastName").send_keys("Raj")
driver.find_element_by_id("text-password").send_keys("isgreat1")
driver.find_element_by_id("text-password2").send_keys("isgreat1")
driver.save_screenshot('C:/Users/KumarGaurav/Desktop/screenTest/screen3.png')
driver.find_element_by_id("btn-create").click()
driver.find_element_by_xpath("//option[@value='In which city were you born? Enter city name only.']").click()
driver.find_element_by_id("text-answer1").send_keys("Bangalore")
driver.find_element_by_xpath("//option[@value='What is the first name of your oldest niece?']").click()
driver.find_element_by_id("text-answer2").send_keys("Gaurav")
driver.save_screenshot('C:/Users/KumarGaurav/Desktop/screenTest/screen4.png')
driver.find_element_by_id("btn-create").click()
assert driver.find_element_by_id("titleDiv").text == "Registration Confirmation"
driver.save_screenshot('C:/Users/KumarGaurav/Desktop/screenTest/screen5.png')
print("Registration is successful")

#Login

driver.execute_script("window.open('');")
driver.switch_to.window(driver.window_handles[2])
# driver.get("https://myaccount..com/registration")
driver.get("https://link..com/#/signin ")
driver.find_element_by_id("username").send_keys(value)
driver.find_element_by_id("password").send_keys("isgreat1")
driver.save_screenshot('C:/Users/KumarGaurav/Desktop/screenTest/screen6.png')
driver.find_element_by_id("signin_button").click()
driver.save_screenshot('C:/Users/KumarGaurav/Desktop/screenTest/screen7.png')
signout= driver.find_element_by_xpath("//a[text()='SIGN OUT']").text
assert "SIGN OUT" == signout
print("Login Successfully")
driver.find_element_by_xpath("//a[text()='SIGN OUT']").click()
print("Sign out Success")

driver.get("https://myaccount..com/forgotpassword/default.aspx?cc=US&lc=en&app=ct&pgm=HL&ru=https%3A%2F%2Flink..com&tgt=mobile")

while driver.current_url != "https://myaccount..com/forgotpassword/forgotpassword-questions":
    driver.find_element_by_id("text-firstName").send_keys(value)
    captchaEle = driver.find_element_by_id("CaptchaImage")
    captchaEle.screenshot('C:/Users/KumarGaurav/Desktop/screenTest/captcha12.png')
    #driver.save_screenshot('C:/Users/KumarGaurav/Desktop/screenTest/screen8.png')
    str1, a = OCR.ocrconvertor()
    time.sleep(5)
    print('{} {}'.format(str1, a))
    driver.find_element_by_id("CaptchaInputText").send_keys(str1)
    time.sleep(2)
    driver.save_screenshot('C:/Users/KumarGaurav/Desktop/screenTest/screen9.png')
    driver.find_element_by_id("btn-next").click()


driver.find_element_by_id("AnswerOne").send_keys("Bangalore")
driver.find_element_by_id("AnswerTwo").send_keys("Gaurav")
time.sleep(1)
driver.save_screenshot('C:/Users/KumarGaurav/Desktop/screenTest/screen10.png')
driver.find_element_by_id("btn-next").click()
driver.find_element_by_id("text-password").send_keys("isgreat2022")
driver.find_element_by_id("text-password2").send_keys("isgreat2022")
driver.save_screenshot('C:/Users/KumarGaurav/Desktop/screenTest/screen11.png')
time.sleep(1)
driver.find_element_by_id("btn-next").click()
text1 = driver.find_element_by_id("titleDiv").text
print(text1)
assert text1 == "CONFIRMATION"
driver.save_screenshot('C:/Users/KumarGaurav/Desktop/screenTest/screen12.png')

doc = docx.Document()

doc.add_heading('CIAM Maintenance', 0)

# Image in its native size
doc.add_heading('Email Verification :', 3)
doc.add_picture('C:/Users/KumarGaurav/Desktop/screenTest/screen1.png',width=Inches(7), height=Inches(4))
doc.add_heading('Email Received  :', 3)
doc.add_picture('C:/Users/KumarGaurav/Desktop/screenTest/screen2.png',width=Inches(7), height=Inches(4))
doc.add_heading('Registration Step 2 :', 3)
doc.add_picture('C:/Users/KumarGaurav/Desktop/screenTest/screen3.png',width=Inches(7), height=Inches(4))
doc.add_heading('Registration Step 3 :', 3)
doc.add_picture('C:/Users/KumarGaurav/Desktop/screenTest/screen4.png',width=Inches(7), height=Inches(4))
doc.add_heading('Registration Completed :', 3)
doc.add_picture('C:/Users/KumarGaurav/Desktop/screenTest/screen5.png',width=Inches(7), height=Inches(4))
doc.add_heading('Sign in :', 3)
doc.add_picture('C:/Users/KumarGaurav/Desktop/screenTest/screen6.png',width=Inches(7), height=Inches(4))
doc.add_heading('Signed in Successfully :', 3)
doc.add_picture('C:/Users/KumarGaurav/Desktop/screenTest/screen7.png',width=Inches(7), height=Inches(4))
doc.add_heading('Signed in Successfully :', 3)
doc.add_picture('C:/Users/KumarGaurav/Desktop/screenTest/screen9.png',width=Inches(7), height=Inches(4))
doc.add_heading('Security question :', 3)
doc.add_picture('C:/Users/KumarGaurav/Desktop/screenTest/screen10.png',width=Inches(7), height=Inches(4))
doc.add_heading('Confirm password :', 3)
doc.add_picture('C:/Users/KumarGaurav/Desktop/screenTest/screen11.png',width=Inches(7), height=Inches(4))
doc.add_heading('Password Changed Successfully :', 3)
doc.add_picture('C:/Users/KumarGaurav/Desktop/screenTest/screen12.png',width=Inches(7), height=Inches(4))
doc.save('C:/Users/KumarGaurav/Desktop/screenTest/Hidas_Sowbhagya.docx')

